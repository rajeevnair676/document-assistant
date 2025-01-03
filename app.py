import os
import pickle
import time
import logging
import streamlit as st
from typing import List
from dotenv import load_dotenv
from docx import Document
import pdfplumber
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.vectorstores import FAISS
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain_groq import ChatGroq
from langchain.schema import Document as LangchainDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyMuPDFLoader
from langchain.prompts import PromptTemplate


logging.basicConfig(
        filename="app.log",
        encoding="utf-8",
        filemode="a",
        format="{asctime} - {levelname} - {message}",
        style="{",
        datefmt="%Y-%m-%d %H:%M",
        level=logging.INFO
 )

import torch
device = torch.device("cuda" if torch.cuda.is_available() else "cpu")
print(device)

# Set up Streamlit configuration
st.set_page_config(page_icon="📄", layout="wide", page_title="Doc & PDF Chatbot with Context")

def icon(emoji: str):
    """Display an emoji as a page icon."""
    st.write(
        f'<span style="font-size: 78px; line-height: 1">{emoji}</span>',
        unsafe_allow_html=True,
    )

icon("🤖")
st.subheader("Document Assistant", divider="rainbow", anchor=False)

groq_api_key = st.sidebar.text_input("Groq API Key", type="password")
# load_dotenv()

VECTORSTORE_PATH = "./vectorstore.pkl"

# Session state initialization
if "messages" not in st.session_state:
    st.session_state.messages = []

if "qa_chain" not in st.session_state:
    st.session_state.qa_chain = None

if "vectorstore_initialized" not in st.session_state:
    st.session_state.vectorstore_initialized = False

if "vectorstore" not in st.session_state:
    st.session_state.vectorstore = None

def pdf_loader(pdf_path):
    loader = PyMuPDFLoader(pdf_path)
    document = loader.load()
    return document

def parse_word_document(uploaded_file) -> List[LangchainDocument]:
    """Parse a Word document into LangChain Document objects."""
    doc = Document(uploaded_file)
    text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
    return split_doc_docx(text)

def parse_pdf_document(uploaded_file) -> List[LangchainDocument]:
    """Parse a PDF document into LangChain Document objects."""
    text = ""
    # try:
    # with pdfplumber.open(uploaded_file) as pdf:
    #     for page in pdf.pages:
    #         text += page.extract_text() or ""
    temp_file = "./temp.pdf"
    with open(temp_file, "wb") as file:
        file.write(uploaded_file.getvalue())
        file_name = uploaded_file.name

    document = pdf_loader(temp_file)
    
    # return split_recursive(document)
    return split_doc_pymupdf(document)
    # except Exception as e:
    #     st.error(f"Failed to process the PDF document: {e}", icon="🚨")
    #     return []

def split_doc_docx(text, max_chunk_size=500) -> List[LangchainDocument]:
    """Split text into manageable chunks."""
    words = text.split()
    chunks = [" ".join(words[i:i + max_chunk_size]) for i in range(0, len(words), max_chunk_size)]
    return [LangchainDocument(page_content=chunk) for chunk in chunks]

def split_doc_pymupdf(documents,max_chunk_size=500):
    texts,metadatas=[],[]
    for doc in documents:
        texts.append(doc.page_content)
        metadatas.append(doc.metadata)
    text = ' '.join(texts)
    words = text.split()
    chunks = [" ".join(words[i:i + max_chunk_size]) for i in range(0, len(words), max_chunk_size)]
    return [LangchainDocument(page_content=chunk) for chunk in chunks]

def split_recursive(text,chunk_size=500):
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size,
                                        chunk_overlap=100)
    return text_splitter.split_documents(text)

def initialize_vectorstore(documents: List[LangchainDocument]):
    """Initialize a FAISS vectorstore with precomputed embeddings."""
    if not st.session_state.vectorstore_initialized:
        with st.spinner("Generating embeddings and initializing vectorstore. This may take a moment..."):
            time.sleep(0.5)
            embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2",
                                               model_kwargs = {'device': device})
            vectorstore = FAISS.from_documents(documents, embeddings)
            st.session_state.vectorstore = vectorstore
            st.session_state.vectorstore_initialized = True
            save_vectorstore(vectorstore)

def initialize_qa_chain():
    """Initialize a conversational QA chain using the vectorstore."""
    if st.session_state.vectorstore and st.session_state.qa_chain is None:
        retriever = st.session_state.vectorstore.as_retriever(search_kwargs={"k": 2})
        memory = ConversationBufferMemory(memory_key="chat_history", return_messages=True)
        # Define your system instruction
        # system_instruction = "The assistant should provide precise and to point answers, explanation only where it is required."

        # # Define your template with the system instruction
        # template = (
        #     f"{system_instruction} "
        #     "Combine the chat history and follow up question into "
        #     "a standalone question. Chat History: {chat_history}"
        #     "Follow up question: {question}"
        # )

        # # Create the prompt template
        # condense_question_prompt = PromptTemplate.from_template(template)
        st.session_state.qa_chain = ConversationalRetrievalChain.from_llm(
            ChatGroq(temperature=0, model_name="llama3-8b-8192",api_key=groq_api_key),
            retriever=retriever,
            memory=memory,
            # condense_question_prompt=condense_question_prompt
        )
        st.success("All set! You can chat with me now!", icon="✅")

def save_vectorstore(vectorstore):
    """Save the vectorstore to disk."""
    with open(VECTORSTORE_PATH, "wb") as f:
        pickle.dump(vectorstore, f)

# File uploader for Word and PDF documents
if not groq_api_key:
    st.warning("Please enter your Groq API key to proceed!", icon="⚠")
else:
    uploaded_file = st.sidebar.file_uploader("Upload a Word or PDF document (.docx, .pdf)", type=["docx", "pdf"])

    print("SESSION STATE: ",st.session_state.keys())
    if uploaded_file and not st.session_state.vectorstore_initialized:
        if os.path.exists(VECTORSTORE_PATH):
            os.remove(VECTORSTORE_PATH)

        with st.spinner("Processing uploaded document..."):
            time.sleep(0.5)
            if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                documents = parse_word_document(uploaded_file)
            elif uploaded_file.type == "application/pdf":
                documents = parse_pdf_document(uploaded_file)
            else:
                st.error("Unsupported file type. Please upload a .docx or .pdf file.", icon="🚨")
                documents = []

        if documents:
            initialize_vectorstore(documents)
            initialize_qa_chain()
            temp_file_path = "./temp.pdf"
            if os.path.exists(temp_file_path):
                os.remove(temp_file_path)
                logging.info(f"Deleted the file {temp_file_path}")

if prompt := st.chat_input("Enter your question here..."):
    # Append the user's question to the chat history
    st.session_state.messages.append({"role": "user", "content": prompt})

    # Display all previous messages
    for message in st.session_state.messages[:-1]:  # Exclude the current question being added
        role = "user" if message["role"] == "user" else "assistant"
        avatar = "👨‍💻" if role == "user" else "🤖"
        with st.chat_message(role, avatar=avatar):
            st.markdown(message["content"])

    # Display the current user's question
    with st.chat_message("user", avatar="👨‍💻"):
        st.markdown(prompt)

    # Placeholder for the assistant's response
    assistant_message = None
    with st.chat_message("assistant", avatar="🤖"):
        assistant_message = st.empty()  # Placeholder for "Generating response..."

    try:
        if st.session_state.qa_chain:
            # Add spinner while generating response
            with st.spinner("Generating response..."):
                response = st.session_state.qa_chain({"question": prompt, "chat_history": st.session_state.messages})
                assistant_response = response.get("answer", "I'm sorry, I couldn't find an answer.")

                # Append assistant response to the chat history
                st.session_state.messages.append({"role": "assistant", "content": assistant_response})

                # Update the placeholder with the actual assistant response
                assistant_message.markdown(assistant_response)
        else:
            st.warning("Please upload a document to initialize the QA system and make sure you have entered the API key. "
                       "You can get an api key by signing in into https://console.groq.com/")
    except Exception as e:
        assistant_message.markdown(f"Error during QA processing: {e}")